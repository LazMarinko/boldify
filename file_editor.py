import os
from docx.shared import RGBColor
from docx import Document
from typing import List
import json

def boldifier(path: str) -> str:
    black: RGBColor = RGBColor(0, 0, 0)
    dark_blue:RGBColor = RGBColor(0, 50, 155)
    bold_color: str
    try:
        program_dir: str = os.getenv("APPDATA")
        full_path: str = program_dir + "\\Boldify\\config.json"
        with open(full_path, "r") as config_file:
            config = json.load(config_file)
            bold_color = config.get("BoldColor")
    except Exception as e:
        print(e)
    doc: Document = Document(path)
    for para in doc.paragraphs:
        lines: List[str] = para.text.split('\n')
        para.clear()
        for line in lines:
            words: List[str] = line.split(" ")
            for word in words:
                word_half1: str = ""
                word_half2: str = ""
                for index, char in enumerate(word):
                    if index < len(word) / 2:
                        word_half1 += char
                    else:
                        word_half2 += char

                run1 = para.add_run(word_half1)
                run1.bold = True
                if bold_color == "Black":
                    run1.font.color.rgb = black
                if bold_color == "Dark Blue":
                    run1.font.color.rgb = dark_blue
                run2 = para.add_run(word_half2 + " ")

    path_split = path.split("\\")
    new_path: str = ""
    for string in path_split:
        if string.__contains__(".docx"):
            new_path += string
    current_save_path: str = ""
    program_dir: str = os.getenv("APPDATA")
    full_path: str = program_dir + "\\Boldify\\config.json"
    try:
        with open(full_path, "r") as config_file:
            config_data: dict = json.load(config_file)
            current_save_path = config_data.get("SavePath")
    except Exception as e:
        print(e)


    doc.save(current_save_path + "\\boldified " + new_path)
    return current_save_path + "\\boldified " + new_path